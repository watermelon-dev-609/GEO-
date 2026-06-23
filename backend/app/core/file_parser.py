"""文件导入解析 — 支持Word(.docx)、PDF、纯文本文件的内容提取"""

from __future__ import annotations
import logging
from io import BytesIO

logger = logging.getLogger(__name__)

# 支持的文件类型
SUPPORTED_TYPES = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".pdf": "application/pdf",
}

# 最大文件大小（字节）
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def parse_file_content(filename: str, file_bytes: bytes) -> dict:
    """根据文件名后缀解析文件内容

    Args:
        filename: 文件名（含后缀，用于判断类型）
        file_bytes: 文件二进制内容

    Returns:
        {
            "text": str,           # 提取的文本内容
            "word_count": int,     # 字数
            "file_type": str,      # 文件类型
            "parse_method": str,   # 解析方式
            "warnings": [str],     # 解析警告
        }
    """
    import os
    ext = os.path.splitext(filename)[1].lower()

    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError(f"文件过大，最大支持 10MB，当前文件 {len(file_bytes) / 1024 / 1024:.1f}MB")

    if ext in (".txt", ".md"):
        return _parse_text(file_bytes, ext)
    elif ext == ".docx":
        return _parse_docx(file_bytes)
    elif ext == ".pdf":
        return _parse_pdf(file_bytes)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，支持: {', '.join(SUPPORTED_TYPES.keys())}")


def _parse_text(file_bytes: bytes, ext: str) -> dict:
    """解析纯文本/Markdown文件"""
    encodings = ["utf-8", "gbk", "gb2312", "utf-16"]
    text = None
    used_encoding = "utf-8"

    for enc in encodings:
        try:
            text = file_bytes.decode(enc)
            used_encoding = enc
            break
        except (UnicodeDecodeError, LookupError):
            continue

    if text is None:
        raise ValueError("无法识别文件编码，请保存为 UTF-8 格式后重试")

    return {
        "text": text.strip(),
        "word_count": len(text),
        "file_type": ext,
        "parse_method": f"文本解析 ({used_encoding})",
        "warnings": [],
    }


def _parse_docx(file_bytes: bytes) -> dict:
    """解析Word文档（.docx）"""
    warnings = []
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            warnings.append("Word文档中未找到文本内容（可能为扫描件图片），请复制文字后粘贴导入")

        text = "\n\n".join(paragraphs)
        return {
            "text": text,
            "word_count": len(text),
            "file_type": ".docx",
            "parse_method": "python-docx 解析",
            "warnings": warnings,
        }
    except ImportError:
        raise ValueError("Word解析库未安装，请运行: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Word文档解析失败: {str(e)}，请确认文件未损坏且为.docx格式")


def _parse_pdf(file_bytes: bytes) -> dict:
    """解析PDF文件"""
    warnings = []
    text_parts = []

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(BytesIO(file_bytes))

        if reader.is_encrypted:
            warnings.append("PDF文件已加密，尝试解密...")
            try:
                reader.decrypt("")
            except Exception:
                raise ValueError("PDF文件已加密且无法自动解密，请先解除密码保护后重试")

        total_pages = len(reader.pages)
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
            except Exception as e:
                warnings.append(f"第{i+1}页文本提取失败: {e}")

        if not text_parts:
            # 尝试用备用方法
            try:
                from io import StringIO
                # PyPDF2 sometimes needs a different extraction approach
                for i, page in enumerate(reader.pages):
                    content = page.extract_text()
                    if content:
                        text_parts.append(content)
            except Exception:
                pass

        text = "\n\n".join(text_parts)

        if not text.strip():
            raise ValueError(
                "PDF中未提取到文字内容。可能原因：\n"
                "1. PDF为扫描图片格式（请使用OCR工具提取文字后粘贴导入）\n"
                "2. PDF为纯图片文件\n"
                "3. 文件已损坏"
            )

        return {
            "text": text,
            "word_count": len(text),
            "file_type": ".pdf",
            "parse_method": f"PyPDF2 解析 ({total_pages}页)",
            "warnings": warnings,
        }
    except ImportError:
        raise ValueError("PDF解析库未安装，请运行: pip install PyPDF2")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"PDF解析失败: {str(e)}，请确认文件未损坏且为PDF格式")
